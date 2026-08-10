"""CV 파일 텍스트 추출.

지원 형식:
  .txt / .md  -> 의존성 없음
  .pdf        -> pypdf         (extras: cvtool[pdf])
  .docx       -> python-docx   (extras: cvtool[docx])

폐쇄망 주의: pdf/docx 라이브러리는 필요할 때만 lazy import 한다.
없으면 그 형식에서만 명확한 에러를 내고, txt 경로는 아무 의존성 없이 동작한다.

참고(수정 금지): 기존 RAG /opt/data-gov/app/ingestion/parsers/ 에 PDF/docx/이미지
추출이 더 정교하게 구현돼 있음. 이미지(스캔 CV) OCR 이 필요해지면 그쪽을 참고.
"""

from __future__ import annotations

from pathlib import Path


class UnsupportedFormat(ValueError):
    pass


def extract_text(path: str | Path) -> str:
    """파일 확장자에 맞춰 순수 텍스트를 추출한다."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(p)
    suffix = p.suffix.lower()
    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        return _extract_pdf(p)
    if suffix == ".docx":
        return _extract_docx(p)
    raise UnsupportedFormat(f"지원하지 않는 형식: {suffix} (지원: .txt .md .pdf .docx)")


def _extract_pdf(p: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - 환경 의존
        raise ImportError(
            "PDF 파싱에는 pypdf 가 필요합니다. `pip install cvtool[pdf]` "
            "(폐쇄망이면 서버에 pypdf 설치 여부를 먼저 확인하세요)."
        ) from exc
    reader = PdfReader(str(p))
    parts = []
    for page in reader.pages:
        # 기본 추출은 2단 편집·표·탭 정렬을 평문으로 뭉개서 CV 를 망가뜨린다.
        # layout 모드는 원본의 공간 배치를 살려준다 (pypdf 4.0+).
        try:
            text = page.extract_text(extraction_mode="layout")
        except Exception:  # noqa: BLE001
            # 구버전 pypdf 는 extraction_mode 를 모르고(TypeError),
            # 일부 PDF 는 layout 모드에서 실패한다. 그때는 기본 모드로.
            try:
                text = page.extract_text()
            except Exception:  # noqa: BLE001 - 한 페이지 실패로 전체를 버리지 않는다
                text = ""
        parts.append(text or "")
    return "\n".join(parts).strip()


def _extract_docx(p: Path) -> str:
    """워드 문서에서 텍스트를 뽑는다.

    ⚠️ document.paragraphs 는 **표 안의 글자를 포함하지 않는다.**
    한국 이력서는 표로 된 것이 아주 흔해서, 문단만 읽으면 이름·연락처·학력이
    통째로 사라진다(실제로 그런 버그가 있었다).
    그래서 문단과 표를 **문서에 나온 순서대로** 함께 읽는다.
    머리글·바닥글에도 연락처가 들어가는 경우가 있어 같이 훑는다.
    """
    try:
        import docx  # python-docx
        from docx.document import Document as _Doc
        from docx.oxml.ns import qn
        from docx.table import Table, _Cell
        from docx.text.paragraph import Paragraph
    except ImportError as exc:  # pragma: no cover - 환경 의존
        raise ImportError(
            "docx 파싱에는 python-docx 가 필요합니다. `pip install cvtool[docx]` "
            "(폐쇄망이면 서버에 python-docx 설치 여부를 먼저 확인하세요)."
        ) from exc

    def blocks(parent):
        """문단과 표를 문서 순서대로 돌려준다."""
        if isinstance(parent, _Doc):
            element = parent.element.body
        elif isinstance(parent, _Cell):
            element = parent._tc
        else:  # pragma: no cover - 방어적
            return
        for child in element.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent)
            elif child.tag == qn("w:tbl"):
                yield Table(child, parent)

    def render(parent) -> list[str]:
        out: list[str] = []
        for block in blocks(parent):
            if isinstance(block, Paragraph):
                if block.text.strip():
                    out.append(block.text)
            else:  # Table
                for row in block.rows:
                    cells = []
                    for cell in row.cells:
                        # 셀 안에 또 표가 있을 수 있다
                        cells.append(" ".join(render(cell)))
                    line = " | ".join(c.strip() for c in cells if c.strip())
                    if line:
                        out.append(line)
        return out

    document = docx.Document(str(p))
    parts = render(document)

    # 머리글·바닥글 (연락처가 여기 있는 이력서가 있다)
    for section in document.sections:
        for area in (section.header, section.footer):
            for para in area.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

    return "\n".join(parts).strip()
